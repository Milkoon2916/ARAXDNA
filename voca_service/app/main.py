import json
import re
import uuid
from pathlib import Path

from fastapi import FastAPI, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from google import genai
from google.genai import types as genai_types
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = BASE_DIR / "app" / "templates"
STATIC_DIR = BASE_DIR / "static"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# PDF는 화면(result.html/worksheet.html)과 동일한 핑크 브랜드 스타일의 HTML을
# WeasyPrint로 그대로 인쇄하는 방식이다 (한글 폰트는 Dockerfile의 fonts-noto-cjk로 해결).
_pdf_env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))

DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
DEFAULT_GEMINI_MODEL = "gemini-2.5-flash"

app = FastAPI(title="ARA Vocab AI")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.exception_handler(HTTPException)
def friendly_error_page(request: Request, exc: HTTPException):
    return HTMLResponse(
        status_code=exc.status_code,
        content=f"""<!DOCTYPE html>
<html lang="ko"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>오류 - ARA Vocab AI</title>
<link rel="stylesheet" href="static/style.css"></head>
<body><div class="container">
<header class="page-header">
  <div class="eyebrow">Error</div>
  <h1>문제가 발생했어요</h1>
  <p>{exc.detail}</p>
</header>
<a class="link-button" href="./">← 처음으로 돌아가기</a>
</div></body></html>""",
    )


# ---------------------------------------------------------------------------
# LLM 호출 (사용자가 입력한 provider / api_key / model 만 사용, 서버에 저장하지 않음)
# ---------------------------------------------------------------------------
def clean_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def call_llm(provider: str, api_key: str, model: str, system_prompt: str, user_prompt: str) -> str:
    """provider별로 사용자 API 키를 사용해 JSON 문자열 응답을 받아온다.
    api_key는 이 함수 호출 동안만 메모리에 존재하고 어디에도 저장되지 않는다."""
    if not api_key or not api_key.strip():
        raise HTTPException(status_code=400, detail="API 키를 입력해주세요.")

    provider = (provider or "gemini").strip().lower()

    if provider == "openai":
        client = OpenAI(api_key=api_key.strip())
        try:
            response = client.chat.completions.create(
                model=(model or DEFAULT_OPENAI_MODEL).strip(),
                temperature=0.2,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"OpenAI 요청 실패: {e}")
        return response.choices[0].message.content

    if provider == "gemini":
        client = genai.Client(api_key=api_key.strip())
        try:
            response = client.models.generate_content(
                model=(model or DEFAULT_GEMINI_MODEL).strip(),
                contents=user_prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    response_mime_type="application/json",
                    temperature=0.2,
                ),
            )
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Gemini 요청 실패: {e}")
        return response.text

    raise HTTPException(status_code=400, detail=f"알 수 없는 provider입니다: {provider}")


def analyze_passage(passage: str, level: str, count: int, focus: str,
                     provider: str, api_key: str, model: str) -> dict:
    prompt = f"""
You are an expert English teacher, vocabulary curriculum designer, and Korean EFL assessment writer.

Analyze the English passage below for {level} students.
Select exactly {count} high-value vocabulary items.

Primary focus: {focus}

Selection rules:
1. Prioritize words that are important for understanding the passage's topic, logic, argument, or narrative.
2. Prefer vocabulary with strong educational value and useful synonym/antonym expansion.
3. Avoid words that are too basic for the stated level unless the contextual meaning is especially important.
4. The meaning must reflect the meaning in this passage, not merely the first dictionary meaning.
5. The context_sentence must be copied exactly from the passage. Never invent an example sentence.
6. Give 2-4 useful synonyms and 1-3 useful antonyms when possible.
7. Include useful derivatives and collocations when appropriate.
8. Keep Korean explanations concise and classroom-friendly.
9. Rank importance from 1 to 5.
10. Difficulty should be one of: 초급, 중급, 고등, 수능.

Return ONLY valid JSON in exactly this structure:
{{
  "title": "short English title",
  "level": "{level}",
  "vocabulary": [
    {{
      "word": "string",
      "part_of_speech": "noun/verb/adjective/adverb/etc.",
      "meaning": "Korean meaning in context",
      "synonyms": ["string", "string"],
      "antonyms": ["string", "string"],
      "context_sentence": "exact sentence or sentence fragment copied from the passage",
      "context_meaning": "short Korean explanation",
      "importance": 1,
      "difficulty": "초급/중급/고등/수능",
      "derivatives": ["string"],
      "collocations": ["string"]
    }}
  ]
}}

Passage:
{passage}
"""
    raw = call_llm(provider, api_key, model, "Return valid JSON only. Do not add markdown.", prompt)
    return clean_json(raw)


def generate_worksheet(data: dict, worksheet_type: str, count: int, difficulty: str,
                        provider: str, api_key: str, model: str) -> dict:
    vocab_json = json.dumps(data.get("vocabulary", []), ensure_ascii=False)
    prompt = f"""
You are an expert Korean EFL test writer.

Create a vocabulary worksheet from the following vocabulary list.

Worksheet type: {worksheet_type}
Number of questions: {count}
Difficulty: {difficulty}

Question types may include:
- 영어 → 한국어 뜻
- 한국어 뜻 → 영어
- synonym multiple choice
- antonym multiple choice
- context cloze
- word order
- mixed

Important:
- Use only the vocabulary and information provided.
- Do not create impossible or ambiguous questions.
- For multiple choice, create exactly 4 choices.
- The answer must be unambiguous.
- Do not show the answer in the question text.

Return ONLY valid JSON:
{{
  "title": "Vocabulary Test",
  "questions": [
    {{
      "number": 1,
      "type": "meaning/synonym/antonym/cloze/order",
      "question": "question text",
      "answer": "correct answer",
      "choices": ["choice1", "choice2", "choice3", "choice4"]
    }}
  ]
}}

Vocabulary:
{vocab_json}
"""
    raw = call_llm(provider, api_key, model, "Return valid JSON only.", prompt)
    return clean_json(raw)


# ---------------------------------------------------------------------------
# DOCX / PDF 생성 (기존 로직과 동일, 입력 데이터에만 의존)
# ---------------------------------------------------------------------------
def add_docx_header(doc, title, subtitle=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(subtitle).italic = True


def create_vocabulary_docx(data: dict, output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_docx_header(doc, "ARA VOCABULARY", f"{data.get('title', '')} · {data.get('level', '')}")

    for i, item in enumerate(data.get("vocabulary", []), 1):
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        rows = [
            ("단어", item.get("word", "")),
            ("품사", item.get("part_of_speech", "")),
            ("뜻", item.get("meaning", "")),
            ("유의어", ", ".join(item.get("synonyms", []))),
            ("반의어", ", ".join(item.get("antonyms", []))),
            ("지문 속 예문", item.get("context_sentence", "")),
            ("문맥상 의미", item.get("context_meaning", "")),
            ("파생어", ", ".join(item.get("derivatives", []))),
            ("연어/표현", ", ".join(item.get("collocations", []))),
        ]

        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        doc.add_paragraph()

    doc.save(output_path)


def create_worksheet_docx(data: dict, worksheet: dict, output_path: Path, answer_key=False):
    doc = Document()
    add_docx_header(doc, "ARA VOCABULARY TEST", data.get("title", ""))

    p = doc.add_paragraph("Name: ______________________________    Date: ________________")
    p.paragraph_format.space_after = Pt(14)

    for q in worksheet.get("questions", []):
        p = doc.add_paragraph()
        p.add_run(f"{q.get('number')}. ").bold = True

        if answer_key:
            p.add_run(f"{q.get('question', '')}  →  {q.get('answer', '')}")
        else:
            p.add_run(q.get("question", ""))
            choices = q.get("choices", [])
            if choices:
                doc.add_paragraph("    " + "   ".join(
                    f"{chr(65+i)}. {choice}" for i, choice in enumerate(choices)
                ))

        doc.add_paragraph("")

    doc.save(output_path)


def create_vocabulary_pdf(data: dict, output_path: Path):
    """화면에 보이는 word-card와 동일한 핑크 브랜드 스타일로 PDF를 렌더링한다."""
    template = _pdf_env.get_template("vocabulary_pdf.html.j2")
    html_str = template.render(
        title=data.get("title", ""),
        level=data.get("level", ""),
        vocabulary=data.get("vocabulary", []),
    )
    HTML(string=html_str).write_pdf(str(output_path))


def create_worksheet_pdf(data: dict, worksheet: dict, output_path: Path, answer_key=False):
    """화면 단어시험지와 동일한 스타일로 PDF를 렌더링한다. answer_key=True면 정답지 스타일."""
    template = _pdf_env.get_template("worksheet_pdf.html.j2")
    html_str = template.render(
        title=worksheet.get("title", data.get("title", "")),
        questions=worksheet.get("questions", []),
        answer_key=answer_key,
    )
    HTML(string=html_str).write_pdf(str(output_path))


# ---------------------------------------------------------------------------
# 라우트
# ---------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def home():
    return (TEMPLATE_DIR / "index.html").read_text(encoding="utf-8")


@app.post("/analyze", response_class=HTMLResponse)
def analyze(
    provider: str = Form("gemini"),
    api_key: str = Form(""),
    model: str = Form(""),
    passage: str = Form(""),
    level: str = Form("고등학교 1학년"),
    count: int = Form(10),
    focus: str = Form("지문 이해와 시험 대비"),
):
    if not passage.strip():
        raise HTTPException(status_code=400, detail="분석할 지문을 입력해주세요.")
    data = analyze_passage(passage, level, count, focus, provider, api_key, model)
    html = (TEMPLATE_DIR / "result.html").read_text(encoding="utf-8")

    cards = []
    for i, item in enumerate(data.get("vocabulary", []), 1):
        cards.append(f"""
        <article class="word-card">
            <div class="word-header">
                <div>
                    <span class="number">{i}</span>
                    <h2>{item.get('word', '')}</h2>
                </div>
                <span class="badge">{item.get('difficulty', '')}</span>
            </div>
            <div class="definition"><b>{item.get('meaning', '')}</b> · {item.get('part_of_speech', '')}</div>
            <div class="grid">
                <div><span>유의어</span><p>{', '.join(item.get('synonyms', []))}</p></div>
                <div><span>반의어</span><p>{', '.join(item.get('antonyms', []))}</p></div>
                <div><span>파생어</span><p>{', '.join(item.get('derivatives', []))}</p></div>
                <div><span>연어/표현</span><p>{', '.join(item.get('collocations', []))}</p></div>
            </div>
            <div class="sentence"><b>지문 속 예문</b><br>{item.get('context_sentence', '')}</div>
            <p class="context"><b>문맥상 의미:</b> {item.get('context_meaning', '')}</p>
        </article>
        """)

    payload = json.dumps(data, ensure_ascii=False)
    html = html.replace("{{TITLE}}", data.get("title", "Vocabulary"))
    html = html.replace("{{LEVEL}}", data.get("level", ""))
    html = html.replace("{{VOCABULARY}}", "\n".join(cards))
    html = html.replace("{{DATA}}", payload)
    html = html.replace("{{PROVIDER}}", provider)
    html = html.replace("{{API_KEY}}", api_key)
    html = html.replace("{{MODEL}}", model)
    return html


@app.post("/generate-worksheet", response_class=HTMLResponse)
def generate_worksheet_route(
    provider: str = Form("gemini"),
    api_key: str = Form(""),
    model: str = Form(""),
    data: str = Form(""),
    worksheet_type: str = Form("혼합형"),
    count: int = Form(20),
    difficulty: str = Form("중상"),
):
    if not data.strip():
        raise HTTPException(status_code=400, detail="단어 데이터가 없습니다. 지문 분석을 다시 진행해주세요.")
    parsed = json.loads(data)
    worksheet = generate_worksheet(parsed, worksheet_type, count, difficulty, provider, api_key, model)
    html = (TEMPLATE_DIR / "worksheet.html").read_text(encoding="utf-8")
    questions = []

    for q in worksheet.get("questions", []):
        choices = q.get("choices", [])
        choice_html = ""
        if choices:
            choice_items = "".join(
                f"<span>{chr(65+i)}. {c}</span>" for i, c in enumerate(choices)
            )
            choice_html = f"<div class='choices'>{choice_items}</div>"
        questions.append(
            f"<div class='question'><b>{q.get('number')}.</b> {q.get('question', '')}{choice_html}</div>"
        )

    html = html.replace("{{TITLE}}", worksheet.get("title", "Vocabulary Test"))
    html = html.replace("{{QUESTIONS}}", "\n".join(questions))
    html = html.replace("{{DATA}}", json.dumps(parsed, ensure_ascii=False))
    html = html.replace("{{WORKSHEET}}", json.dumps(worksheet, ensure_ascii=False))
    return html


@app.post("/download/vocabulary/docx")
def download_vocabulary_docx(data: str = Form(...)):
    parsed = json.loads(data)
    path = OUTPUT_DIR / f"ara_vocabulary_{uuid.uuid4().hex}.docx"
    create_vocabulary_docx(parsed, path)
    return FileResponse(path, filename="ara_vocabulary.docx",
                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/download/vocabulary/pdf")
def download_vocabulary_pdf(data: str = Form(...)):
    parsed = json.loads(data)
    path = OUTPUT_DIR / f"ara_vocabulary_{uuid.uuid4().hex}.pdf"
    create_vocabulary_pdf(parsed, path)
    return FileResponse(path, filename="ara_vocabulary.pdf", media_type="application/pdf")


@app.post("/download/worksheet/docx")
def download_worksheet_docx(data: str = Form(...), worksheet: str = Form(...), answer_key: bool = Form(False)):
    parsed = json.loads(data)
    worksheet_data = json.loads(worksheet)
    path = OUTPUT_DIR / f"ara_worksheet_{uuid.uuid4().hex}.docx"
    create_worksheet_docx(parsed, worksheet_data, path, answer_key)
    filename = "ara_worksheet_answer_key.docx" if answer_key else "ara_worksheet.docx"
    return FileResponse(path, filename=filename,
                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/download/worksheet/pdf")
def download_worksheet_pdf(data: str = Form(...), worksheet: str = Form(...), answer_key: bool = Form(False)):
    parsed = json.loads(data)
    worksheet_data = json.loads(worksheet)
    path = OUTPUT_DIR / f"ara_worksheet_{uuid.uuid4().hex}.pdf"
    create_worksheet_pdf(parsed, worksheet_data, path, answer_key)
    filename = "ara_worksheet_answer_key.pdf" if answer_key else "ara_worksheet.pdf"
    return FileResponse(path, filename=filename, media_type="application/pdf")
